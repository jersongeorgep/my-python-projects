# lib/iso_templates.py
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Mm


def set_page_layout(doc: Document):
    """A4 page size + 0.5 inch margins."""
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        margin = Inches(0.5)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def add_iso_cover_page(
    doc: Document,
    company_name: str,
    document_title: str,
    project_name: str,
    project_code: str,
    version: str,
    prepared_by: str,
    approved_by: str,
    date_str: str,
    logo_image=None,
):
    cover_section = doc.sections[0]

    if logo_image:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        try:
            run_logo.add_picture(logo_image, width=Inches(2))
        except Exception:
            pass

    doc.add_paragraph("\n")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(document_title)
    run_title.bold = True
    run_title.font.size = Pt(24)

    doc.add_paragraph("\n")

    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_company = p_company.add_run(company_name)
    run_company.font.size = Pt(16)
    run_company.bold = True

    doc.add_paragraph("\n\n")

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    data = [
        ("Project Name", project_name),
        ("Project Code", project_code),
        ("Document Version", version),
        ("Prepared By", prepared_by),
        ("Approved By", approved_by),
        ("Date", date_str),
    ]

    for row, (label, value) in zip(table.rows, data):
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph("\n\n")

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run("ISO/IEC 27001:2022")
    run_footer.font.size = Pt(10)
    run_footer.italic = True

    doc.add_page_break()


def add_custom_header(doc: Document, company_name, document_name, project_code, date_str, logo_image=None):
    section = doc.sections[0]
    header = section.header

    for para in header.paragraphs:
        p = para._p
        p.getparent().remove(p)

    table = header.add_table(rows=1, cols=3, width=Inches(8.27))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.0)
    table.columns[2].width = Inches(2.0)

    cell_left = table.cell(0, 0)
    cell_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left = p_left.add_run()

    if logo_image:
        try:
            run_left.add_picture(logo_image, width=Inches(1.2))
        except Exception:
            pass

    cell_center = table.cell(0, 1)
    cell_center.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p_center = cell_center.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_center = p_center.add_run(company_name + "\n")
    run_center.bold = True
    run_center.font.size = Pt(14)

    run_doc = p_center.add_run(document_name)
    run_doc.bold = True
    run_doc.font.size = Pt(12)

    cell_right = table.cell(0, 2)
    cell_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run("Project Code: " + project_code)
    p_right.add_run("\nDate: " + date_str)


def add_custom_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer

    if footer.paragraphs:
        p1 = footer.paragraphs[0]
    else:
        p1 = footer.add_paragraph()

    p1.text = "Security Classification: Internal Use Only"
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    fldPage = OxmlElement("w:fldSimple")
    fldPage.set(qn("w:instr"), "PAGE")
    p2._p.append(fldPage)

    p2.add_run(" of ")

    fldNumPages = OxmlElement("w:fldSimple")
    fldNumPages.set(qn("w:instr"), "NUMPAGES")
    p2._p.append(fldNumPages)


def generate_iso27001_policy(
    title: str,
    company: str,
    prepared_by: str,
    sections: list,
    logo_image=None,
) -> BytesIO:
    doc = Document()

    if logo_image:
        try:
            doc.add_picture(logo_image, width=Inches(1.5))
        except Exception:
            pass

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.cell(0, 0).text = "Company"
    meta.cell(0, 1).text = company
    meta.cell(1, 0).text = "Policy"
    meta.cell(1, 1).text = title
    meta.cell(2, 0).text = "Prepared By"
    meta.cell(2, 1).text = prepared_by
    meta.cell(3, 0).text = "Date"
    meta.cell(3, 1).text = datetime.today().strftime("%Y-%m-%d")

    for sec_title, content in sections:
        doc.add_heading(sec_title, level=2)
        for line in str(content or "").replace("<br>", "\n").split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    footer = doc.sections[0].footer
    if footer.paragraphs:
        footer.paragraphs[0].text = f"ISO/IEC 27001:2022 – {company}"
    else:
        footer.add_paragraph(f"ISO/IEC 27001:2022 – {company}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
