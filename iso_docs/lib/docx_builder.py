# lib/docx_builder.py
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Inches

from html2docx import html2docx

from .html_parser import clean_html
from .iso_templates import (
    set_page_layout,
    add_iso_cover_page,
    add_custom_header,
    add_custom_footer,
    generate_iso27001_policy as _generate_iso27001_policy,
)


def _append_html(doc: Document, html: str | None):
    """
    Convert HTML → real DOCX content (bold, italics, etc.)
    So LibreOffice / Word renders correctly (no raw <p> tags).
    """
    html = clean_html(html)
    if not html.strip():
        return
    # html2docx directly mutates the Document
    html2docx(html, doc)


def add_table(doc: Document, df):
    if df is None or df.empty:
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if val is None else str(val)


def create_iso_docx_iso_template(
    title: str,
    project_name: str,
    project_code: str,
    company_name: str,
    prepared_by: str,
    reviewed_by: str = "ISMS Core Team",
    approved_by: str = "CISO",
    version: str = "1.0",
    content_sections=None,
    logo_image=None,
    doc_date: str | None = None,
) -> BytesIO:
    """
    ISO-style template with cover page, header, footer, document control etc.
    Uses html2docx so rich text is respected.
    """
    doc = Document()
    set_page_layout(doc)

    if doc_date is None:
        doc_date = datetime.today().strftime("%Y-%m-%d")

    # Cover page
    add_iso_cover_page(
        doc,
        company_name=company_name,
        document_title=title,
        project_name=project_name,
        project_code=project_code,
        version=version,
        prepared_by=prepared_by,
        approved_by=approved_by,
        date_str=doc_date,
        logo_image=logo_image,
    )

    # Header
    add_custom_header(
        doc,
        company_name=company_name,
        document_name=title,
        project_code=project_code,
        date_str=doc_date,
        logo_image=logo_image,
    )

    # Title
    heading = doc.add_heading(f"{company_name}\n{title}", level=1)
    heading.alignment = 1  # center

    # Disclaimer
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        f"This document contains confidential information for {company_name}. "
        "Unauthorized disclosure or distribution is prohibited."
    )

    # Document Control
    doc.add_heading("Document Control", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Document Name"
    table.cell(0, 1).text = title
    table.cell(1, 0).text = "Project Code"
    table.cell(1, 1).text = project_code
    table.cell(2, 0).text = "Version"
    table.cell(2, 1).text = version
    table.cell(3, 0).text = "Date"
    table.cell(3, 1).text = doc_date

    # Authorization
    doc.add_heading("Authorization", level=2)
    auth = doc.add_table(rows=3, cols=4)
    auth.style = "Table Grid"
    auth.cell(0, 0).text = "Prepared By"
    auth.cell(0, 1).text = prepared_by
    auth.cell(1, 0).text = "Reviewed By"
    auth.cell(1, 1).text = reviewed_by
    auth.cell(2, 0).text = "Approved By"
    auth.cell(2, 1).text = approved_by

    # Version History
    doc.add_heading("Version History", level=2)
    vh = doc.add_table(rows=2, cols=4)
    vh.style = "Table Grid"
    vh.cell(0, 0).text = "Version"
    vh.cell(0, 1).text = "Changes"
    vh.cell(0, 2).text = "Date"
    vh.cell(0, 3).text = "Prepared By"
    vh.cell(1, 0).text = version
    vh.cell(1, 1).text = "Initial Release"
    vh.cell(1, 2).text = doc_date
    vh.cell(1, 3).text = prepared_by

    # Main content
    if content_sections:
        for sec_title, sec_html in content_sections:
            doc.add_heading(sec_title, level=2)
            _append_html(doc, sec_html)

    # Footer
    add_custom_footer(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_iso_docx(
    title: str,
    sections=None,
    tables=None,
    image=None,
    image_caption: str | None = None,
    meta: dict | None = None,
) -> BytesIO:
    """
    Generic DOCX builder using html2docx for sections.
    """
    from .docx_builder import add_table as _add_table  # self import fn only, safe

    doc = Document()

    if meta is None:
        meta = {}

    company_name = meta.get("company_name", "")
    project_name = meta.get("project_name", "")
    project_id = meta.get("project_id", "")
    doc_version = meta.get("doc_version", "")
    prepared_by = meta.get("prepared_by", "")
    approved_by = meta.get("approved_by", "")
    logo_image = meta.get("logo_image", None)

    if logo_image:
        try:
            doc.add_picture(logo_image, width=Inches(1.5))
        except Exception:
            pass

    doc.add_heading(title, level=1)

    meta_table = doc.add_table(rows=0, cols=2)

    def add_meta_row(label, value):
        if not (label or value):
            return
        row_cells = meta_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    today_str = datetime.today().strftime("%Y-%m-%d")

    add_meta_row("Company", company_name)
    add_meta_row("Project", project_name)
    add_meta_row("Project ID", project_id)
    add_meta_row("Document Version", doc_version)
    add_meta_row("Prepared By", prepared_by)
    add_meta_row("Approved By", approved_by)
    add_meta_row("Date", today_str)

    doc.add_paragraph("")

    if sections:
        for sec_title, sec_html in sections:
            if sec_title:
                doc.add_heading(sec_title, level=2)
            _append_html(doc, sec_html)

    if tables:
        for tbl_title, df in tables:
            if df is None or df.empty:
                continue
            if tbl_title:
                doc.add_heading(tbl_title, level=2)
            _add_table(doc, df)

    if image:
        if image_caption:
            doc.add_heading(image_caption, level=2)
        doc.add_picture(image, width=Inches(6))

    try:
        section = doc.sections[0]
        footer = section.footer
        footer_text = f"Confidential – {company_name}".strip(" –")
        if footer.paragraphs:
            footer.paragraphs[0].text = footer_text
        else:
            footer.add_paragraph(footer_text)
    except Exception:
        pass

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_iso27001_policy(*args, **kwargs) -> BytesIO:
    """
    Thin wrapper so app imports from docx_builder.
    """
    return _generate_iso27001_policy(*args, **kwargs)
