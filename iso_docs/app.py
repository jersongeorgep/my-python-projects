import os
import tempfile
from io import BytesIO
from zipfile import ZipFile
from datetime import datetime, date

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sqlalchemy import create_engine, text
import pymysql  # noqa – required for mysql+pymysql

# ===================== CONFIG =====================

# MySQL connection – adjust if needed
ENGINE = create_engine(
    "mysql+pymysql://root:@localhost/iso_data?charset=utf8mb4",
    echo=False,
    future=True,
)

# Optional PDF export
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

VALID_USERS = {
    "admin": "admin123",  # change this for production
}

# ===================== DB HELPERS =====================

def db_get_projects_list() -> pd.DataFrame:
    query = "SELECT id, name FROM projects ORDER BY id DESC"
    with ENGINE.connect() as conn:
        return pd.read_sql(query, conn)


def db_get_project(project_id: int):
    query = "SELECT * FROM projects WHERE id = :id"
    with ENGINE.connect() as conn:
        df = pd.read_sql(text(query), conn, params={"id": project_id})
    return df.to_dict(orient="records")[0] if not df.empty else None


def db_get_df(table: str, project_id: int) -> pd.DataFrame:
    query = f"SELECT * FROM {table} WHERE project_id = :pid"
    with ENGINE.connect() as conn:
        return pd.read_sql(text(query), conn, params={"pid": project_id})


def db_save_df(table: str, project_id: int, df: pd.DataFrame, id_col: str = "id"):
    """
    Replace all rows for project_id in a child table.
    """
    with ENGINE.begin() as conn:
        conn.execute(text(f"DELETE FROM {table} WHERE project_id = :pid"), {"pid": project_id})
        if df is not None and not df.empty:
            df = df.copy()
            df["project_id"] = project_id
            if id_col in df.columns:
                df = df.drop(columns=[id_col])
            df.to_sql(table, conn, if_exists="append", index=False)


def db_get_section(project_id: int, key: str) -> str:
    query = """
        SELECT content FROM project_sections
        WHERE project_id = :pid AND section_key = :key
        LIMIT 1
    """
    with ENGINE.connect() as conn:
        df = pd.read_sql(text(query), conn, params={"pid": project_id, "key": key})
    return df["content"].iloc[0] if not df.empty else ""


def db_save_section(project_id: int, key: str, content: str):
    query = """
        INSERT INTO project_sections (project_id, section_key, content)
        VALUES (:pid, :key, :content)
        ON DUPLICATE KEY UPDATE content = VALUES(content)
    """
    with ENGINE.begin() as conn:
        conn.execute(text(query), {"pid": project_id, "key": key, "content": content})


def upsert_project(
    project_name: str,
    client_name: str,
    project_code: str,
    start_date: date,
    end_date: date,
    overview: str,
    scope_in_scope: str,
    scope_out_scope: str,
) -> int:
    """
    If project with same name or project_code exists:
        delete it and all its children, then insert a fresh row.
    Return new project_id.
    """
    with ENGINE.begin() as conn:
        existing = conn.execute(
            text(
                """
                SELECT id FROM projects
                WHERE name = :name OR project_code = :code
                LIMIT 1
                """
            ),
            {"name": project_name, "code": project_code},
        ).fetchone()

        if existing:
            old_id = existing[0]
            # delete children first
            conn.execute(text("DELETE FROM project_team_members WHERE project_id=:pid"), {"pid": old_id})
            conn.execute(text("DELETE FROM project_tasks WHERE project_id=:pid"), {"pid": old_id})
            conn.execute(text("DELETE FROM project_risks WHERE project_id=:pid"), {"pid": old_id})
            conn.execute(text("DELETE FROM project_test_cases WHERE project_id=:pid"), {"pid": old_id})
            conn.execute(text("DELETE FROM project_apis WHERE project_id=:pid"), {"pid": old_id})
            conn.execute(text("DELETE FROM project_change_requests WHERE project_id=:pid"), {"pid": old_id})
            conn.execute(text("DELETE FROM project_sections WHERE project_id=:pid"), {"pid": old_id})
            conn.execute(text("DELETE FROM projects WHERE id=:pid"), {"pid": old_id})

        result = conn.execute(
            text(
                """
                INSERT INTO projects
                (name, client_name, project_code, start_date, end_date,
                 overview, scope_in_scope, scope_out_scope)
                VALUES
                (:name, :client, :code, :start, :end, :overview, :in_scope, :out_scope)
                """
            ),
            {
                "name": project_name,
                "client": client_name,
                "code": project_code,
                "start": start_date,
                "end": end_date,
                "overview": overview,
                "in_scope": scope_in_scope,
                "out_scope": scope_out_scope,
            },
        )
        return result.lastrowid

# ===================== SIMPLE TEXT EDITOR =====================

def rich_editor(label: str, key: str, initial: str = "") -> str:
    """
    Simple, stable text editor: just a text area.
    No HTML, no Quill, no delta => no 'List argument must consist only of dictionaries'.
    """
    return st.text_area(label, value=initial or "", key=key, height=160)

# ===================== DOCX HELPERS =====================

def set_page_layout(doc: Document):
    """A4 page, 0.5 inch margins."""
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        margin = Inches(0.5)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def add_cover_page(
    doc: Document,
    company_name: str,
    document_title: str,
    project_name: str,
    project_code: str,
    version: str,
    prepared_by: str,
    approved_by: str,
    date_str: str,
    logo_image: BytesIO | None = None,
):
    """
    Fully centered cover page, no header/footer.
    Then start a NEW SECTION (next page) for contents.
    """
    # Logo
    if logo_image:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_logo = p_logo.add_run()
        try:
            r_logo.add_picture(logo_image, width=Inches(2))
        except Exception:
            pass

    doc.add_paragraph()  # spacing

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(document_title)
    r_title.bold = True
    r_title.font.size = Pt(26)

    doc.add_paragraph()

    # Company
    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_company = p_company.add_run(company_name)
    r_company.bold = True
    r_company.font.size = Pt(18)

    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    data = [
        ("Project Name", project_name),
        ("Project Code", project_code),
        ("Document Version", version),
        ("Prepared By", prepared_by),
        ("Approved By", approved_by),
        ("Date", date_str),
    ]
    for (label, value), row in zip(data, table.rows):
        row.cells[0].text = label
        row.cells[1].text = value or ""

    doc.add_paragraph()

    # Footer note on cover
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run("ISO/IEC 27001:2022")
    r_footer.italic = True
    r_footer.font.size = Pt(10)

    # New section for rest of document
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_table_of_contents(doc: Document):
    """
    Insert a TOC field. User must update fields in Word/LibreOffice.
    """
    p_heading = doc.add_paragraph()
    p_heading.style = "Heading 1"
    p_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_heading.add_run("Table of Contents")

    p_toc = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p_toc._p.append(fld)


def add_header_footer(
    doc,
    company_name,
    document_name,
    project_code,
    date_str,
    logo_image=None,
):
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    for i, section in enumerate(doc.sections):
        if i == 0:
            continue  # skip cover page

        # ---------------- HEADER ----------------
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing
        for p in header.paragraphs:
            p._element.getparent().remove(p._element)

        # FIXED: header table requires width
        table = header.add_table(rows=1, cols=3, width=Inches(6.5))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # LEFT CELL
        c1 = table.cell(0, 0)
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p1.add_run()

        if logo_image:
            try:
                run1.add_picture(logo_image, width=Inches(1))
            except:
                pass
        p1.add_run("\nProject Code: " + project_code)

        # CENTER CELL
        c2 = table.cell(0, 1)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"{company_name}\n{document_name}")
        r2.bold = True

        # RIGHT CELL
        c3 = table.cell(0, 2)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.add_run("Date: " + date_str)

        # ---------------- FOOTER ----------------
        footer = section.footer
        footer.is_linked_to_previous = False

        for p in footer.paragraphs:
            p._element.getparent().remove(p._element)

        p_left = footer.add_paragraph()
        p_left.text = "Internal Use Only"
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p_page = footer.add_paragraph()
        p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fld_page = OxmlElement("w:fldSimple")
        fld_page.set(qn("w:instr"), "PAGE")
        p_page._p.append(fld_page)

        p_page.add_run(" of ")

        fld_total = OxmlElement("w:fldSimple")
        fld_total.set(qn("w:instr"), "NUMPAGES")
        p_page._p.append(fld_total)


def add_df_table(doc: Document, df: pd.DataFrame, title: str | None = None):
    """
    Insert heading + table from DataFrame.
    """
    if df is None or df.empty:
        return

    if title:
        h = doc.add_heading(title, level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        hdr_cells[idx].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for idx, val in enumerate(row):
            row_cells[idx].text = "" if pd.isna(val) else str(val)


def create_iso_docx_iso_template(
    title: str,
    project_name: str,
    project_code: str,
    company_name: str,
    prepared_by: str,
    reviewed_by: str = "ISMS Core Team",
    approved_by: str = "CISO",
    version: str = "1.0",
    content_sections: list[tuple[str, str]] | None = None,
    table_sections: list[tuple[str, pd.DataFrame]] | None = None,
    logo_image: BytesIO | None = None,
) -> BytesIO:
    """
    Build a DOCX with:
      - Cover page
      - TOC
      - Header+Footer (with page numbers)
      - Document control, authorization, version history
      - Plain text sections
      - Optional DataFrame tables
    """
    doc = Document()
    set_page_layout(doc)

    today_str = datetime.today().strftime("%Y-%m-%d")

    # Cover
    add_cover_page(
        doc,
        company_name=company_name,
        document_title=title,
        project_name=project_name,
        project_code=project_code,
        version=version,
        prepared_by=prepared_by,
        approved_by=approved_by,
        date_str=today_str,
        logo_image=logo_image,
    )

    # TOC
    add_table_of_contents(doc)

    # Header & footer for all non-cover sections
    add_header_footer(
        doc,
        company_name=company_name,
        document_name=title,
        project_code=project_code,
        date_str=today_str,
        logo_image=logo_image,
    )

    # Main heading
    heading = doc.add_heading(f"{company_name}\n{title}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Disclaimer
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        f"This document contains confidential information for {company_name}. "
        "Unauthorized disclosure or distribution is prohibited."
    )

    # Document control
    doc.add_heading("Document Control", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Document Name"
    table.cell(0, 1).text = title
    table.cell(1, 0).text = "Project Code"
    table.cell(1, 1).text = project_code
    table.cell(2, 0).text = "Version"
    table.cell(2, 1).text = version
    table.cell(3, 0).text = "Date"
    table.cell(3, 1).text = today_str

    # Authorization
    doc.add_heading("Authorization", level=2)
    auth = doc.add_table(rows=3, cols=4)
    auth.style = "Table Grid"
    auth.cell(0, 0).text = "Prepared By"
    auth.cell(0, 1).text = prepared_by
    auth.cell(1, 0).text = "Reviewed By"
    auth.cell(1, 1).text = reviewed_by
    auth.cell(2, 0).text = "Approved By"
    auth.cell(2, 1).text = approved_by

    # Version history
    doc.add_heading("Version History", level=2)
    vh = doc.add_table(rows=2, cols=4)
    vh.style = "Table Grid"
    vh.cell(0, 0).text = "Version"
    vh.cell(0, 1).text = "Changes"
    vh.cell(0, 2).text = "Date"
    vh.cell(0, 3).text = "Prepared By"
    vh.cell(1, 0).text = version
    vh.cell(1, 1).text = "Initial Release"
    vh.cell(1, 2).text = today_str
    vh.cell(1, 3).text = prepared_by

    # Content sections
    if content_sections:
        for sec_title, sec_text in content_sections:
            if not sec_title and not sec_text:
                continue
            doc.add_heading(sec_title, level=2)
            if sec_text:
                for line in str(sec_text).splitlines():
                    doc.add_paragraph(line)

    # DataFrame tables
    if table_sections:
        for tbl_title, df in table_sections:
            add_df_table(doc, df, title=tbl_title)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ===================== GANTT CHART =====================

def create_gantt_chart(tasks_df: pd.DataFrame) -> BytesIO | None:
    if tasks_df is None or tasks_df.empty:
        return None

    df = tasks_df.copy()
    required_cols = {"Task", "Start", "End"}
    if not required_cols.issubset(df.columns):
        return None

    df["Start"] = pd.to_datetime(df["Start"], errors="coerce")
    df["End"] = pd.to_datetime(df["End"], errors="coerce")
    df = df.dropna(subset=["Start", "End", "Task"])

    if df.empty:
        return None

    df = df.sort_values("Start")

    fig, ax = plt.subplots(figsize=(8, 4))
    for _, row in df.iterrows():
        duration = (row["End"] - row["Start"]).days
        if duration <= 0:
            duration = 1
        ax.barh(row["Task"], duration, left=row["Start"])

    ax.set_xlabel("Date")
    ax.set_ylabel("Task")
    ax.set_title("Project Gantt Chart")
    fig.autofmt_xdate()

    img = BytesIO()
    fig.savefig(img, format="png", bbox_inches="tight")
    img.seek(0)
    plt.close(fig)
    return img

# ===================== AUTH & UTIL =====================

def show_login():
    st.title("🔐 ISO Documentation Portal – Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if username in VALID_USERS and VALID_USERS[username] == password:
            st.session_state.authenticated = True
            st.session_state.username = username
            st.success("Logged in successfully")
            st.rerun()
        else:
            st.error("Invalid credentials")


def load_section(project_id, key):
    if project_id is None:
        return ""
    try:
        return db_get_section(project_id, key)
    except Exception:
        return ""


def wkey(name: str, current_project_id):
    pid = current_project_id if current_project_id is not None else "new"
    return f"{name}_{pid}"

# ===================== STREAMLIT APP =====================

st.set_page_config(page_title="ISO 9001 + 27001 Documentation Suite", layout="wide")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    show_login()
    st.stop()

st.sidebar.markdown(f"**Logged in as:** {st.session_state.username}")
if st.sidebar.button("Logout"):
    st.session_state.authenticated = False
    st.rerun()

st.title("📚 ISO/IEC 27001:2022 Documentation Generator")

# ----- Sidebar meta -----
with st.sidebar:
    st.header("Company & Document Meta")
    company_name = st.text_input("Company Name", value="Your Company")
    doc_version = st.text_input("Document Version", value="1.0")
    prepared_by = st.text_input("Prepared By", value=st.session_state.username)
    approved_by = st.text_input("Approved By", value="")

    logo_file = st.file_uploader("Company Logo (PNG/JPG)", type=["png", "jpg", "jpeg"])

    st.markdown("---")
    st.subheader("Project Selection")

    projects_df = db_get_projects_list()
    project_options = ["-- New Project --"]
    id_to_name = {}
    for _, row in projects_df.iterrows():
        label = f"{row['id']} – {row['name']}"
        project_options.append(label)
        id_to_name[label] = row["id"]

    selected_project_label = st.selectbox("Select Project", project_options, index=0)
    current_project_id = None
    current_project_row = None
    if selected_project_label != "-- New Project --":
        current_project_id = id_to_name[selected_project_label]
        current_project_row = db_get_project(current_project_id)

# ----- Project header fields -----
if current_project_row:
    default_project_name = current_project_row.get("name") or ""
    default_client_name = current_project_row.get("client_name") or ""
    default_project_code = current_project_row.get("project_code") or ""
    default_start_date = current_project_row.get("start_date") or datetime.today().date()
    default_end_date = current_project_row.get("end_date") or datetime.today().date()
    default_overview = current_project_row.get("overview") or ""
    default_scope_in_scope = current_project_row.get("scope_in_scope") or ""
    default_scope_out_scope = current_project_row.get("scope_out_scope") or ""
else:
    default_project_name = "Sample Project"
    default_client_name = ""
    default_project_code = "PRJ-001"
    default_start_date = datetime.today().date()
    default_end_date = datetime.today().date()
    default_overview = ""
    default_scope_in_scope = ""
    default_scope_out_scope = ""

col1, col2, col3 = st.columns(3)
with col1:
    project_name = st.text_input("Project Name", value=default_project_name)
with col2:
    client_name = st.text_input("Client / Organization", value=default_client_name)
with col3:
    project_code = st.text_input("Project ID / Code", value=default_project_code)

col4, col5 = st.columns(2)
with col4:
    start_date = st.date_input("Start Date", value=default_start_date)
with col5:
    end_date = st.date_input("End Date", value=default_end_date)

st.markdown("---")

(
    tab_overview,
    tab_scope,
    tab_team,
    tab_req,
    tab_design,
    tab_plan,
    tab_risk,
    tab_security,
    tab_quality,
    tab_tests,
    tab_api,
    tab_deploy,
    tab_meet_change,
    tab_vc_release,
    tab_backup_bcp,
    tab_user_docs,
) = st.tabs(
    [
        "Overview",
        "Scope",
        "Team",
        "SRS (Requirements)",
        "SDS (Design)",
        "Plan & Gantt",
        "Risk Assessment",
        "Security (ISO 27001)",
        "Quality (ISO 9001)",
        "Testing",
        "API Docs",
        "Deployment",
        "Meetings & Changes",
        "Version Control & Release",
        "Backup & Continuity",
        "User Docs",
    ]
)

# ---------- Overview ----------
with tab_overview:
    st.subheader("Project Overview")
    project_overview = rich_editor(
        "High-level overview (business context, objectives, constraints):",
        key=wkey("project_overview_editor", current_project_id),
        initial=default_overview,
    )

# ---------- Scope ----------
with tab_scope:
    st.subheader("Project Scope (ISO 9001 & 27001 Context)")
    scope_in_scope = rich_editor(
        "In-scope items (modules, locations, processes):",
        key=wkey("scope_in_scope_editor", current_project_id),
        initial=default_scope_in_scope,
    )
    scope_out_scope = rich_editor(
        "Out-of-scope items (what is explicitly excluded):",
        key=wkey("scope_out_scope_editor", current_project_id),
        initial=default_scope_out_scope,
    )

# ---------- Team ----------
with tab_team:
    st.subheader("Project Team & Responsibilities")
    if current_project_id:
        team_db = db_get_df("project_team_members", current_project_id)
        if not team_db.empty:
            default_team = team_db[["name", "role", "responsibility"]].rename(
                columns={"name": "Name", "role": "Role", "responsibility": "Responsibility"}
            )
        else:
            default_team = pd.DataFrame(
                [
                    {"Name": "Project Manager", "Role": "PM", "Responsibility": "Overall delivery"},
                    {"Name": "Lead Developer", "Role": "Tech Lead", "Responsibility": "Architecture & development"},
                ]
            )
    else:
        default_team = pd.DataFrame(
            [
                {"Name": "Project Manager", "Role": "PM", "Responsibility": "Overall delivery"},
                {"Name": "Lead Developer", "Role": "Tech Lead", "Responsibility": "Architecture & development"},
            ]
        )

    team_df = st.data_editor(
        default_team,
        num_rows="dynamic",
        width="stretch",
        key=wkey("team_editor", current_project_id),
    )

# ---------- SRS ----------
with tab_req:
    st.subheader("Software Requirements Specification (SRS)")
    srs_intro = rich_editor(
        "Introduction & Purpose:",
        key=wkey("srs_intro_editor", current_project_id),
        initial=load_section(current_project_id, "srs_intro"),
    )
    srs_scope = rich_editor(
        "System Scope (from business perspective):",
        key=wkey("srs_scope_editor", current_project_id),
        initial=load_section(current_project_id, "srs_scope"),
    )
    srs_functional = rich_editor(
        "Functional Requirements:",
        key=wkey("srs_functional_editor", current_project_id),
        initial=load_section(current_project_id, "srs_functional"),
    )
    srs_nonfunctional = rich_editor(
        "Non-functional Requirements (performance, usability, etc.):",
        key=wkey("srs_nonfunctional_editor", current_project_id),
        initial=load_section(current_project_id, "srs_nonfunctional"),
    )
    srs_constraints = rich_editor(
        "Assumptions & Constraints:",
        key=wkey("srs_constraints_editor", current_project_id),
        initial=load_section(current_project_id, "srs_constraints"),
    )

# ---------- SDS ----------
with tab_design:
    st.subheader("Software Design Specification (SDS)")
    sds_architecture = rich_editor(
        "Architecture Overview:",
        key=wkey("sds_architecture_editor", current_project_id),
        initial=load_section(current_project_id, "sds_architecture"),
    )
    sds_modules = rich_editor(
        "Module / Component Design:",
        key=wkey("sds_modules_editor", current_project_id),
        initial=load_section(current_project_id, "sds_modules"),
    )
    sds_db = rich_editor(
        "Database / Data Model Design:",
        key=wkey("sds_db_editor", current_project_id),
        initial=load_section(current_project_id, "sds_db"),
    )
    sds_interfaces = rich_editor(
        "External Interfaces (systems, services):",
        key=wkey("sds_interfaces_editor", current_project_id),
        initial=load_section(current_project_id, "sds_interfaces"),
    )

# ---------- Plan & Gantt ----------
with tab_plan:
    st.subheader("Project Plan & Tasks (Gantt Input)")
    if current_project_id:
        tasks_db = db_get_df("project_tasks", current_project_id)
        if not tasks_db.empty:
            default_tasks = tasks_db[["task", "owner", "start_date", "end_date", "status"]].rename(
                columns={
                    "task": "Task",
                    "owner": "Owner",
                    "start_date": "Start",
                    "end_date": "End",
                    "status": "Status",
                }
            )
        else:
            default_tasks = pd.DataFrame(
                [
                    {"Task": "Requirements", "Owner": "PM", "Start": start_date, "End": start_date, "Status": "Planned"},
                    {"Task": "Development", "Owner": "Dev Team", "Start": start_date, "End": end_date, "Status": "Planned"},
                ]
            )
    else:
        default_tasks = pd.DataFrame(
            [
                {"Task": "Requirements", "Owner": "PM", "Start": start_date, "End": start_date, "Status": "Planned"},
                {"Task": "Development", "Owner": "Dev Team", "Start": start_date, "End": end_date, "Status": "Planned"},
            ]
        )

    tasks_df = st.data_editor(
        default_tasks,
        num_rows="dynamic",
        width="stretch",
        key=wkey("tasks_editor", current_project_id),
    )

    plan_notes = rich_editor(
        "Additional Planning Notes / Milestones:",
        key=wkey("plan_notes_editor", current_project_id),
        initial=load_section(current_project_id, "plan_notes"),
    )

# ---------- Risk Assessment ----------
with tab_risk:
    st.subheader("Risk Assessment (ISO 27001 + 9001)")
    if current_project_id:
        risks_db = db_get_df("project_risks", current_project_id)
        if not risks_db.empty:
            default_risks = risks_db[
                ["risk_id", "description", "likelihood", "impact", "treatment"]
            ].rename(
                columns={
                    "risk_id": "Risk ID",
                    "description": "Description",
                    "likelihood": "Likelihood",
                    "impact": "Impact",
                    "treatment": "Treatment",
                }
            )
        else:
            default_risks = pd.DataFrame(
                [
                    {"Risk ID": "R-001", "Description": "Data breach", "Likelihood": "High", "Impact": "High",
                     "Treatment": "Encryption, access control"},
                    {"Risk ID": "R-002", "Description": "Downtime", "Likelihood": "Medium", "Impact": "High",
                     "Treatment": "Redundancy, backup"},
                ]
            )
    else:
        default_risks = pd.DataFrame(
            [
                {"Risk ID": "R-001", "Description": "Data breach", "Likelihood": "High", "Impact": "High",
                 "Treatment": "Encryption, access control"},
                {"Risk ID": "R-002", "Description": "Downtime", "Likelihood": "Medium", "Impact": "High",
                 "Treatment": "Redundancy, backup"},
            ]
        )

    risks_df = st.data_editor(
        default_risks,
        num_rows="dynamic",
        width="stretch",
        key=wkey("risk_editor", current_project_id),
    )

    risk_methodology = rich_editor(
        "Risk Assessment Methodology:",
        key=wkey("risk_methodology_editor", current_project_id),
        initial=load_section(current_project_id, "risk_methodology"),
    )
    risk_treatment_overview = rich_editor(
        "Risk Treatment Plan Overview:",
        key=wkey("risk_treatment_overview_editor", current_project_id),
        initial=load_section(current_project_id, "risk_treatment_overview"),
    )

# ---------- Security (ISO 27001) ----------
with tab_security:
    st.subheader("Information Security (ISO 27001)")
    sec_policy = rich_editor(
        "Information Security Policy:",
        key=wkey("sec_policy_editor", current_project_id),
        initial=load_section(current_project_id, "sec_policy"),
    )
    sec_asset_inventory = rich_editor(
        "Asset Inventory & Ownership:",
        key=wkey("sec_asset_inventory_editor", current_project_id),
        initial=load_section(current_project_id, "sec_asset_inventory"),
    )
    sec_access_control = rich_editor(
        "Access Control & User Management:",
        key=wkey("sec_access_control_editor", current_project_id),
        initial=load_section(current_project_id, "sec_access_control"),
    )
    sec_data_class = rich_editor(
        "Data Classification & Handling:",
        key=wkey("sec_data_class_editor", current_project_id),
        initial=load_section(current_project_id, "sec_data_class"),
    )
    sec_encryption = rich_editor(
        "Encryption & Key Management:",
        key=wkey("sec_encryption_editor", current_project_id),
        initial=load_section(current_project_id, "sec_encryption"),
    )
    sec_incident_mgmt = rich_editor(
        "Incident Management Process:",
        key=wkey("sec_incident_mgmt_editor", current_project_id),
        initial=load_section(current_project_id, "sec_incident_mgmt"),
    )
    sec_logging = rich_editor(
        "Logging, Monitoring & Audit Trail:",
        key=wkey("sec_logging_editor", current_project_id),
        initial=load_section(current_project_id, "sec_logging"),
    )

# ---------- Quality (ISO 9001) ----------
with tab_quality:
    st.subheader("Quality Management (ISO 9001)")
    q_quality_policy = rich_editor(
        "Quality Policy & Objectives:",
        key=wkey("q_quality_policy_editor", current_project_id),
        initial=load_section(current_project_id, "q_quality_policy"),
    )
    q_process_overview = rich_editor(
        "Process Approach & Workflow:",
        key=wkey("q_process_overview_editor", current_project_id),
        initial=load_section(current_project_id, "q_process_overview"),
    )
    q_quality_plan = rich_editor(
        "Quality Plan for this Project:",
        key=wkey("q_quality_plan_editor", current_project_id),
        initial=load_section(current_project_id, "q_quality_plan"),
    )
    q_change_mgmt = rich_editor(
        "Change Management Process:",
        key=wkey("q_change_mgmt_editor", current_project_id),
        initial=load_section(current_project_id, "q_change_mgmt"),
    )
    q_document_control = rich_editor(
        "Document & Record Control:",
        key=wkey("q_document_control_editor", current_project_id),
        initial=load_section(current_project_id, "q_document_control"),
    )
    q_ncr = rich_editor(
        "Non-conformity & Corrective Action Process:",
        key=wkey("q_ncr_editor", current_project_id),
        initial=load_section(current_project_id, "q_ncr"),
    )

# ---------- Testing ----------
with tab_tests:
    st.subheader("Test Strategy & Test Cases")
    test_strategy = rich_editor(
        "Test Strategy (types of testing, environments, entry/exit criteria):",
        key=wkey("test_strategy_editor", current_project_id),
        initial=load_section(current_project_id, "test_strategy"),
    )

    if current_project_id:
        tests_db = db_get_df("project_test_cases", current_project_id)
        if not tests_db.empty:
            default_tests = tests_db[
                ["test_case_id", "description", "preconditions", "steps", "expected", "status"]
            ].rename(
                columns={
                    "test_case_id": "Test Case ID",
                    "description": "Description",
                    "preconditions": "Preconditions",
                    "steps": "Steps",
                    "expected": "Expected",
                    "status": "Status",
                }
            )
        else:
            default_tests = pd.DataFrame(
                [
                    {"Test Case ID": "TC-001", "Description": "Valid login", "Preconditions": "User exists",
                     "Steps": "Enter valid creds", "Expected": "Dashboard", "Status": "Not Run"},
                    {"Test Case ID": "TC-002", "Description": "Invalid login", "Preconditions": "",
                     "Steps": "Enter wrong password", "Expected": "Error message", "Status": "Not Run"},
                ]
            )
    else:
        default_tests = pd.DataFrame(
            [
                {"Test Case ID": "TC-001", "Description": "Valid login", "Preconditions": "User exists",
                 "Steps": "Enter valid creds", "Expected": "Dashboard", "Status": "Not Run"},
                {"Test Case ID": "TC-002", "Description": "Invalid login", "Preconditions": "",
                 "Steps": "Enter wrong password", "Expected": "Error message", "Status": "Not Run"},
            ]
        )

    tests_df = st.data_editor(
        default_tests,
        num_rows="dynamic",
        width="stretch",
        key=wkey("tests_editor", current_project_id),
    )

# ---------- API Docs ----------
with tab_api:
    st.subheader("API Documentation")
    api_overview = rich_editor(
        "API Overview (purpose, architecture):",
        key=wkey("api_overview_editor", current_project_id),
        initial=load_section(current_project_id, "api_overview"),
    )

    if current_project_id:
        api_db = db_get_df("project_apis", current_project_id)
        if not api_db.empty:
            default_apis = api_db[
                ["endpoint", "method", "description", "auth", "params_body"]
            ].rename(
                columns={
                    "endpoint": "Endpoint",
                    "method": "Method",
                    "description": "Description",
                    "auth": "Auth",
                    "params_body": "Params/Body",
                }
            )
        else:
            default_apis = pd.DataFrame(
                [
                    {"Endpoint": "/api/login", "Method": "POST", "Description": "Authenticate user",
                     "Auth": "JWT", "Params/Body": "login, password"},
                    {"Endpoint": "/api/projects", "Method": "GET", "Description": "List projects",
                     "Auth": "JWT", "Params/Body": ""},
                ]
            )
    else:
        default_apis = pd.DataFrame(
            [
                {"Endpoint": "/api/login", "Method": "POST", "Description": "Authenticate user",
                 "Auth": "JWT", "Params/Body": "login, password"},
                {"Endpoint": "/api/projects", "Method": "GET", "Description": "List projects",
                 "Auth": "JWT", "Params/Body": ""},
            ]
        )

    api_df = st.data_editor(
        default_apis,
        num_rows="dynamic",
        width="stretch",
        key=wkey("api_editor", current_project_id),
    )

# ---------- Deployment ----------
with tab_deploy:
    st.subheader("Deployment & Rollback Checklist")
    deploy_pre = rich_editor(
        "Pre-deployment Checks:",
        key=wkey("deploy_pre_editor", current_project_id),
        initial=load_section(current_project_id, "deploy_pre"),
    )
    deploy_steps = rich_editor(
        "Deployment Steps:",
        key=wkey("deploy_steps_editor", current_project_id),
        initial=load_section(current_project_id, "deploy_steps"),
    )
    rollback_steps = rich_editor(
        "Rollback Procedure:",
        key=wkey("rollback_steps_editor", current_project_id),
        initial=load_section(current_project_id, "rollback_steps"),
    )
    env_details = rich_editor(
        "Environment Details (DEV/UAT/PROD):",
        key=wkey("env_details_editor", current_project_id),
        initial=load_section(current_project_id, "env_details"),
    )

# ---------- Meetings & Changes ----------
with tab_meet_change:
    st.subheader("Meetings & Change Requests")
    meeting_template = rich_editor(
        "Meeting Minutes Template / Notes:",
        key=wkey("meeting_template_editor", current_project_id),
        initial=load_section(current_project_id, "meeting_template") or
                "Date:\nAttendees:\nAgenda:\nDiscussion Points:\nDecisions:\nAction Items & Owners:\n",
    )

    if current_project_id:
        change_db = db_get_df("project_change_requests", current_project_id)
        if not change_db.empty:
            default_changes = change_db[
                ["cr_id", "title", "raised_by", "impact", "status"]
            ].rename(
                columns={
                    "cr_id": "CR ID",
                    "title": "Title",
                    "raised_by": "Raised By",
                    "impact": "Impact",
                    "status": "Status",
                }
            )
        else:
            default_changes = pd.DataFrame(
                [
                    {"CR ID": "CR-001", "Title": "Change home page layout",
                     "Raised By": "Client", "Impact": "Medium", "Status": "Proposed"},
                ]
            )
    else:
        default_changes = pd.DataFrame(
            [
                {"CR ID": "CR-001", "Title": "Change home page layout",
                 "Raised By": "Client", "Impact": "Medium", "Status": "Proposed"},
            ]
        )

    change_df = st.data_editor(
        default_changes,
        num_rows="dynamic",
        width="stretch",
        key=wkey("change_editor", current_project_id),
    )

# ---------- Version Control & Release ----------
with tab_vc_release:
    st.subheader("Version Control & Release Management")
    git_repo = st.text_input(
        "Git Repository URL",
        value=load_section(current_project_id, "git_repo") or "https://github.com/org/repo",
    )
    git_branch = st.text_input(
        "Main Branch",
        value=load_section(current_project_id, "git_branch") or "main",
    )
    git_workflow = rich_editor(
        "Branching & Merge Strategy (Git Flow, etc.):",
        key=wkey("git_workflow_editor", current_project_id),
        initial=load_section(current_project_id, "git_workflow"),
    )
    version_history = rich_editor(
        "Version / Release Notes Summary:",
        key=wkey("version_history_editor", current_project_id),
        initial=load_section(current_project_id, "version_history"),
    )

# ---------- Backup & Continuity ----------
with tab_backup_bcp:
    st.subheader("Backup, Recovery & Business Continuity")
    backup_policy = rich_editor(
        "Backup Policy (frequency, scope, tools):",
        key=wkey("backup_policy_editor", current_project_id),
        initial=load_section(current_project_id, "backup_policy"),
    )
    restore_procedure = rich_editor(
        "Restore Procedure:",
        key=wkey("restore_procedure_editor", current_project_id),
        initial=load_section(current_project_id, "restore_procedure"),
    )
    bcp_plan = rich_editor(
        "Business Continuity & Disaster Recovery Plan:",
        key=wkey("bcp_plan_editor", current_project_id),
        initial=load_section(current_project_id, "bcp_plan"),
    )

# ---------- User Docs ----------
with tab_user_docs:
    st.subheader("User Documentation")
    user_manual = rich_editor(
        "User Manual (high-level skeleton):",
        key=wkey("user_manual_editor", current_project_id),
        initial=load_section(current_project_id, "user_manual"),
    )
    training_plan = rich_editor(
        "Training & Handover Plan:",
        key=wkey("training_plan_editor", current_project_id),
        initial=load_section(current_project_id, "training_plan"),
    )

st.markdown("---")

# ===================== SAVE + GENERATE DOCS =====================

if st.button("💾 Save to DB & 📦 Generate All ISO Documents (ZIP)"):
    try:
        # 1) Save/Update project
        project_id = upsert_project(
            project_name=project_name,
            client_name=client_name,
            project_code=project_code,
            start_date=start_date,
            end_date=end_date,
            overview=project_overview,
            scope_in_scope=scope_in_scope,
            scope_out_scope=scope_out_scope,
        )

        # 2) Save sections
        section_map = {
            "srs_intro": srs_intro,
            "srs_scope": srs_scope,
            "srs_functional": srs_functional,
            "srs_nonfunctional": srs_nonfunctional,
            "srs_constraints": srs_constraints,
            "sds_architecture": sds_architecture,
            "sds_modules": sds_modules,
            "sds_db": sds_db,
            "sds_interfaces": sds_interfaces,
            "plan_notes": plan_notes,
            "risk_methodology": risk_methodology,
            "risk_treatment_overview": risk_treatment_overview,
            "sec_policy": sec_policy,
            "sec_asset_inventory": sec_asset_inventory,
            "sec_access_control": sec_access_control,
            "sec_data_class": sec_data_class,
            "sec_encryption": sec_encryption,
            "sec_incident_mgmt": sec_incident_mgmt,
            "sec_logging": sec_logging,
            "q_quality_policy": q_quality_policy,
            "q_process_overview": q_process_overview,
            "q_quality_plan": q_quality_plan,
            "q_change_mgmt": q_change_mgmt,
            "q_document_control": q_document_control,
            "q_ncr": q_ncr,
            "test_strategy": test_strategy,
            "api_overview": api_overview,
            "deploy_pre": deploy_pre,
            "deploy_steps": deploy_steps,
            "rollback_steps": rollback_steps,
            "env_details": env_details,
            "meeting_template": meeting_template,
            "git_repo": git_repo,
            "git_branch": git_branch,
            "git_workflow": git_workflow,
            "version_history": version_history,
            "backup_policy": backup_policy,
            "restore_procedure": restore_procedure,
            "bcp_plan": bcp_plan,
            "user_manual": user_manual,
            "training_plan": training_plan,
        }
        for key, content in section_map.items():
            db_save_section(project_id, key, content or "")

        # 3) Save table data
        db_save_df(
            "project_team_members",
            project_id,
            team_df.rename(columns={
                "Name": "name",
                "Role": "role",
                "Responsibility": "responsibility",
            }),
        )

        db_save_df(
            "project_tasks",
            project_id,
            tasks_df.rename(columns={
                "Task": "task",
                "Owner": "owner",
                "Start": "start_date",
                "End": "end_date",
                "Status": "status",
            }),
        )

        db_save_df(
            "project_risks",
            project_id,
            risks_df.rename(columns={
                "Risk ID": "risk_id",
                "Description": "description",
                "Likelihood": "likelihood",
                "Impact": "impact",
                "Treatment": "treatment",
            }),
        )

        db_save_df(
            "project_test_cases",
            project_id,
            tests_df.rename(columns={
                "Test Case ID": "test_case_id",
                "Description": "description",
                "Preconditions": "preconditions",
                "Steps": "steps",
                "Expected": "expected",
                "Status": "status",
            }),
        )

        db_save_df(
            "project_apis",
            project_id,
            api_df.rename(columns={
                "Endpoint": "endpoint",
                "Method": "method",
                "Description": "description",
                "Auth": "auth",
                "Params/Body": "params_body",
            }),
        )

        db_save_df(
            "project_change_requests",
            project_id,
            change_df.rename(columns={
                "CR ID": "cr_id",
                "Title": "title",
                "Raised By": "raised_by",
                "Impact": "impact",
                "Status": "status",
            }),
        )

        # 4) Generate DOCX
        temp_dir = tempfile.mkdtemp()
        folder_name = project_name.replace(" ", "_")
        project_folder = os.path.join(temp_dir, folder_name)
        os.makedirs(project_folder, exist_ok=True)

        logo_bytes = None
        if logo_file is not None:
            logo_bytes = BytesIO(logo_file.read())

        serial = 1

        # Overview
        overview_doc = create_iso_docx_iso_template(
            title="Project Overview",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Purpose", project_overview),
                ("Timeline", f"Start: {start_date}\nEnd: {end_date}"),
                ("Client Details", f"Client: {client_name}"),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Project_Overview.docx"), "wb") as f:
            f.write(overview_doc.read())
        serial += 1

        # Scope
        scope_doc = create_iso_docx_iso_template(
            title="Project Scope",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("In Scope", scope_in_scope),
                ("Out of Scope", scope_out_scope),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Project_Scope.docx"), "wb") as f:
            f.write(scope_doc.read())
        serial += 1

        # Team (Roles & Responsibilities as DOCX table)
        team_doc = create_iso_docx_iso_template(
            title="Team Roles & Responsibilities",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Team Structure", "Project team overview"),
            ],
            table_sections=[
                ("Roles & Responsibilities", team_df[["Name", "Role", "Responsibility"]]),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Team_and_Roles.docx"), "wb") as f:
            f.write(team_doc.read())
        serial += 1

        # SRS
        srs_doc = create_iso_docx_iso_template(
            title="Software Requirements Specification (SRS)",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Introduction", srs_intro),
                ("Scope", srs_scope),
                ("Functional Requirements", srs_functional),
                ("Non-Functional Requirements", srs_nonfunctional),
                ("Constraints", srs_constraints),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_SRS.docx"), "wb") as f:
            f.write(srs_doc.read())
        serial += 1

        # SDS
        sds_doc = create_iso_docx_iso_template(
            title="Software Design Specification (SDS)",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Architecture Overview", sds_architecture),
                ("Modules & Components", sds_modules),
                ("Database Design", sds_db),
                ("System Interfaces", sds_interfaces),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_SDS.docx"), "wb") as f:
            f.write(sds_doc.read())
        serial += 1

        # Plan & Gantt
        gantt_image = create_gantt_chart(tasks_df)
        plan_doc = create_iso_docx_iso_template(
            title="Project Plan & Milestones",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Planning Notes", plan_notes),
            ],
            table_sections=[
                ("Work Breakdown Structure (WBS)", tasks_df[["Task", "Owner", "Start", "End", "Status"]]),
            ],
            logo_image=logo_bytes,
        )
        # (optional: add Gantt image manually later if needed)
        with open(os.path.join(project_folder, f"{serial:03d}_Project_Plan_and_Gantt.docx"), "wb") as f:
            f.write(plan_doc.read())
        serial += 1

        # Risk Assessment
        risk_doc = create_iso_docx_iso_template(
            title="Risk Assessment & Treatment",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Methodology", risk_methodology),
                ("Risk Treatment Overview", risk_treatment_overview),
            ],
            table_sections=[
                ("Risk Register", risks_df[["Risk ID", "Description", "Likelihood", "Impact", "Treatment"]]),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Risk_Assessment.docx"), "wb") as f:
            f.write(risk_doc.read())
        serial += 1

        # Security
        sec_doc = create_iso_docx_iso_template(
            title="Information Security Plan (ISO 27001)",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Security Policy", sec_policy),
                ("Asset Inventory", sec_asset_inventory),
                ("Access Control", sec_access_control),
                ("Data Classification", sec_data_class),
                ("Encryption", sec_encryption),
                ("Incident Management", sec_incident_mgmt),
                ("Logging & Monitoring", sec_logging),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Information_Security.docx"), "wb") as f:
            f.write(sec_doc.read())
        serial += 1

        # Quality
        quality_doc = create_iso_docx_iso_template(
            title="Quality Management Plan (ISO 9001)",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Quality Policy", q_quality_policy),
                ("Process Workflow", q_process_overview),
                ("Quality Plan", q_quality_plan),
                ("Change Management", q_change_mgmt),
                ("Document Control", q_document_control),
                ("Corrective Actions", q_ncr),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Quality_Management.docx"), "wb") as f:
            f.write(quality_doc.read())
        serial += 1

        # Testing
        test_doc = create_iso_docx_iso_template(
            title="Test Strategy & Test Cases",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Test Strategy", test_strategy),
            ],
            table_sections=[
                ("Test Cases", tests_df[["Test Case ID", "Description", "Preconditions", "Steps", "Expected", "Status"]]),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Testing.docx"), "wb") as f:
            f.write(test_doc.read())
        serial += 1

        # API Documentation
        api_doc = create_iso_docx_iso_template(
            title="API Documentation",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("API Overview", api_overview),
            ],
            table_sections=[
                ("Endpoints", api_df[["Endpoint", "Method", "Description", "Auth", "Params/Body"]]),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_API_Documentation.docx"), "wb") as f:
            f.write(api_doc.read())
        serial += 1

        # Deployment
        deploy_doc = create_iso_docx_iso_template(
            title="Deployment & Rollback Plan",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Pre-deployment Checklist", deploy_pre),
                ("Deployment Steps", deploy_steps),
                ("Rollback Procedure", rollback_steps),
                ("Environment Details", env_details),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Deployment_and_Rollback.docx"), "wb") as f:
            f.write(deploy_doc.read())
        serial += 1

        # Meetings
        meetings_doc = create_iso_docx_iso_template(
            title="Meeting Minutes",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Template", meeting_template),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Meeting_Minutes.docx"), "wb") as f:
            f.write(meetings_doc.read())
        serial += 1

        # Change Requests
        change_doc = create_iso_docx_iso_template(
            title="Change Request Log",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            table_sections=[
                ("Change Requests", change_df[["CR ID", "Title", "Raised By", "Impact", "Status"]]),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Change_Request_Log.docx"), "wb") as f:
            f.write(change_doc.read())
        serial += 1

        # Version Control & Release
        vc_doc = create_iso_docx_iso_template(
            title="Version Control & Release Notes",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Repository", f"Repo: {git_repo}\nMain Branch: {git_branch}"),
                ("Workflow", git_workflow),
                ("Release Notes", version_history),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Version_Control_and_Release.docx"), "wb") as f:
            f.write(vc_doc.read())
        serial += 1

        # Backup & BCP
        backup_doc = create_iso_docx_iso_template(
            title="Backup & Business Continuity Plan",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("Backup Policy", backup_policy),
                ("Restore Procedure", restore_procedure),
                ("Business Continuity", bcp_plan),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_Backup_and_BCP.docx"), "wb") as f:
            f.write(backup_doc.read())
        serial += 1

        # User Docs
        user_doc = create_iso_docx_iso_template(
            title="User Documentation",
            project_name=project_name,
            project_code=project_code,
            company_name=company_name,
            prepared_by=prepared_by,
            version=doc_version,
            content_sections=[
                ("User Manual", user_manual),
                ("Training Plan", training_plan),
            ],
            logo_image=logo_bytes,
        )
        with open(os.path.join(project_folder, f"{serial:03d}_User_Documentation.docx"), "wb") as f:
            f.write(user_doc.read())
        serial += 1

        # 5) Optional PDF conversion
        pdf_folder = os.path.join(temp_dir, f"{folder_name}_PDF")
        pdf_generated = False
        if DOCX2PDF_AVAILABLE:
            try:
                os.makedirs(pdf_folder, exist_ok=True)
                for filename in os.listdir(project_folder):
                    if filename.lower().endswith(".docx"):
                        src = os.path.join(project_folder, filename)
                        dst = os.path.join(pdf_folder, filename.replace(".docx", ".pdf"))
                        docx2pdf_convert(src, dst)
                        pdf_generated = True
            except Exception:
                pdf_generated = False

        # 6) Build ZIP
        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, "w") as zipf:
            for filename in os.listdir(project_folder):
                zipf.write(
                    os.path.join(project_folder, filename),
                    arcname=f"{folder_name}/DOCX/{filename}",
                )
            if pdf_generated:
                for filename in os.listdir(pdf_folder):
                    zipf.write(
                        os.path.join(pdf_folder, filename),
                        arcname=f"{folder_name}/PDF/{filename}",
                    )

        zip_buffer.seek(0)

        st.success("Saved to DB and generated all ISO 9001 + 27001 documents successfully!")

        st.download_button(
            label="⬇️ Download All Documents (ZIP)",
            data=zip_buffer,
            file_name=f"{folder_name}_ISO_Documents.zip",
            mime="application/zip",
        )

        if not DOCX2PDF_AVAILABLE:
            st.info("DOCX2PDF not available – only DOCX files included. Install `docx2pdf` + Word/LibreOffice for PDFs.")

    except Exception as e:
        st.error(f"Error: {e}")
